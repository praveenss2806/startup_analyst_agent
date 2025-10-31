import json
import re
import os
import hashlib
import datetime
from collections import Counter
from google.cloud import vision
from google.cloud import storage
import docx
import extract_msg
from google.api_core.client_options import ClientOptions
from google.cloud import documentai_v1beta3 as documentai
from docx import Document
import email
import mimetypes
import tempfile
from typing import Dict, Any, List, Optional, Tuple
from google.adk.tools import ToolContext
import logging


def doc_ingestion_tool(file_path: str, tool_context: ToolContext) -> Dict[str, Any]:
    """
    Intelligently processes documents using Google Cloud Document AI OCR and other text extraction methods.
    Provides comprehensive document analysis including content categorization, key information extraction,
    and detailed metadata analysis. Supports PDF, images, DOCX, MSG, EML, and TXT files from local paths or GCS URLs.

    Args:
        file_path (str): Path to the input file (local path or GCS URL).
        tool_context (ToolContext): The tool context containing state information.
        
    Returns:
        Dict[str, Any]: A comprehensive dictionary containing:
            - status: Success or Failure
            - document_analysis: Detailed analysis results
            - extracted_text: The extracted text content
            - file_metadata: File information and processing details
            - content_analysis: Content categorization and key information
            - quality_metrics: Text quality and confidence scores
    """
    try:
        # Google Cloud Document AI configuration
        PROJECT_ID = os.getenv("GOOGLE_CLOUD_PROJECT", "woven-perigee-476815-m8")
        LOCATION = "us"
        PROCESSOR_ID = "abb1ab40cbac8a9c"
        
        # Check if file_path is a GCS URL or local path
        is_gcs_url = file_path.startswith('gs://') or file_path.startswith('https://storage.googleapis.com/')
        
        if is_gcs_url:
            # Extract file extension from GCS URL
            if file_path.startswith('https://storage.googleapis.com/'):
                # Convert public URL to GCS URI format
                gcs_uri = file_path.replace('https://storage.googleapis.com/', 'gs://')
            else:
                gcs_uri = file_path
            
            # Extract filename from GCS URI
            filename = gcs_uri.split('/')[-1]
            _, file_extension = os.path.splitext(filename)
        else:
            # Local file path
            filename = os.path.basename(file_path)
            _, file_extension = os.path.splitext(file_path)
            
        file_extension = file_extension.lower()
        
        # Initialize analysis results
        extracted_text = ""
        processing_method = ""
        file_size = 0
        processing_time = 0
        
        # Get file size
        if is_gcs_url:
            try:
                bucket_name = gcs_uri.split('/')[2]
                blob_name = '/'.join(gcs_uri.split('/')[3:])
                storage_client = storage.Client(project=PROJECT_ID)
                bucket = storage_client.bucket(bucket_name)
                blob = bucket.blob(blob_name)
                file_size = blob.size
            except:
                file_size = 0
        else:
            try:
                file_size = os.path.getsize(file_path)
            except:
                file_size = 0
        
        start_time = datetime.datetime.now()
        
        # Process file based on extension
        if file_extension in ['.pdf', '.tiff']:
            if is_gcs_url:
                extracted_text = _ocr_pdf_document_gcs(PROJECT_ID, LOCATION, PROCESSOR_ID, gcs_uri)
                processing_method = "Google Cloud Document AI (PDF/TIFF from GCS)"
            else:
                extracted_text = _ocr_pdf_document(PROJECT_ID, LOCATION, PROCESSOR_ID, file_path)
                processing_method = "Google Cloud Document AI (PDF/TIFF)"
            
        elif file_extension in ['.png', '.jpeg', '.jpg']:
            if is_gcs_url:
                extracted_text = _ocr_img_gcs(PROJECT_ID, LOCATION, PROCESSOR_ID, gcs_uri)
                processing_method = "Google Cloud Document AI (Image from GCS)"
            else:
                extracted_text = _ocr_img(PROJECT_ID, LOCATION, PROCESSOR_ID, file_path)
                processing_method = "Google Cloud Document AI (Image)"
            
        elif file_extension == '.docx':
            if is_gcs_url:
                extracted_text = _get_text_from_docx_gcs(gcs_uri)
                processing_method = "python-docx (from GCS)"
            else:
                extracted_text = _get_text_from_docx(file_path)
                processing_method = "python-docx"
            
        elif file_extension == '.eml':
            if is_gcs_url:
                extracted_text = _get_text_from_eml_gcs(gcs_uri)
                processing_method = "Python email library (from GCS)"
            else:
                extracted_text = _get_text_from_eml(file_path)
                processing_method = "Python email library"
            
        elif file_extension == '.msg':
            if is_gcs_url:
                extracted_text = _get_text_from_msg_gcs(gcs_uri)
                processing_method = "extract-msg (from GCS)"
            else:
                extracted_text = _get_text_from_msg(file_path)
                processing_method = "extract-msg"
            
        elif file_extension == '.txt':
            if is_gcs_url:
                extracted_text = _get_text_from_txt_gcs(gcs_uri)
                processing_method = "Direct text reading (from GCS)"
            else:
                extracted_text = _get_text_from_txt(file_path)
                processing_method = "Direct text reading"
            
        else:
            return {
                "status": "failure",
                "error": f"Unsupported file format: {file_extension}",
                "supported_formats": [".pdf", ".tiff", ".png", ".jpg", ".jpeg", ".docx", ".eml", ".msg", ".txt"]
            }
        
        processing_time = (datetime.datetime.now() - start_time).total_seconds()
        
        # Perform comprehensive document analysis
        document_analysis = _analyze_document_content(extracted_text, file_extension, filename)
        
        # Generate file metadata
        file_metadata = _generate_file_metadata(
            filename, file_extension, file_size, processing_method, 
            processing_time, is_gcs_url, file_path
        )
        
        # Analyze content quality and structure
        quality_metrics = _analyze_content_quality(extracted_text, file_extension)
        
        # Extract key information based on document type
        content_analysis = _extract_key_information(extracted_text, file_extension, filename)
        
        # Store extracted data in tool context
        if "startup_information" in tool_context.state:
            tool_context.state["startup_information"] += "\n\n\n\n\n" + extracted_text
        else:
            tool_context.state["startup_information"] = extracted_text
        
        tool_context.state['file_path'] = file_path
        tool_context.state['file_type'] = file_extension
        tool_context.state['processing_method'] = processing_method
        tool_context.state['document_analysis'] = document_analysis
        tool_context.state['file_metadata'] = file_metadata
        tool_context.state['content_analysis'] = content_analysis
        tool_context.state['quality_metrics'] = quality_metrics
        
        return {
            "status": "success",
            "document_analysis": document_analysis,
            "extracted_text": extracted_text,
            "file_metadata": file_metadata,
            "content_analysis": content_analysis,
            "quality_metrics": quality_metrics
        }
        
    except Exception as e:
        logging.error(f"Error in document_ai_ocr_tool: {str(e)}")
        return {
            "status": "failure", 
            "error": str(e),
            "error_type": type(e).__name__,
            "timestamp": datetime.datetime.now().isoformat()
        }


# def _ocr_pdf_document(project_id: str, location: str, processor_id: str, file_path: str) -> str:
#     """
#     Performs OCR on a local PDF file using the Google Cloud Document AI API.
#     """
#     try:
#         client_options = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
#         client = documentai.DocumentProcessorServiceClient(client_options=client_options)
        
#         processor_name = client.processor_path(project_id, location, processor_id)
        
#         with open(file_path, "rb") as document_file:
#             document_content = document_file.read()
        
#         raw_document = documentai.RawDocument(
#             content=document_content,
#             mime_type="application/pdf"
#         )
        
#         request = documentai.ProcessRequest(
#             name=processor_name,
#             raw_document=raw_document,
#             imageless_mode=True
#         )
        
#         response = client.process_document(request=request)
#         document = response.document
        
#         return document.text
        
#     except Exception as e:
#         raise Exception(f"Error processing PDF with Document AI: {e}")
        
def _ocr_pdf_document(project_id: str, location: str, processor_id: str, file_path: str) -> str:
    """
    Performs OCR on a local PDF file using Google Cloud Document AI API.
    Automatically switches to batch processing for large PDFs (>30 pages).
    """
    try:
        client_options = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=client_options)
        processor_name = client.processor_path(project_id, location, processor_id)

        with open(file_path, "rb") as document_file:
            document_content = document_file.read()

        # --- First, use normal sync processing for small docs ---
        try:
            raw_document = documentai.RawDocument(
                content=document_content,
                mime_type="application/pdf"
            )

            request = documentai.ProcessRequest(
                name=processor_name,
                raw_document=raw_document
            )

            response = client.process_document(request=request)
            return response.document.text

        except Exception as sync_err:
            # --- If we hit a page limit or large file, use async batch processing ---
            logging.warning(f"Switching to async batch processing: {sync_err}")

            storage_client = storage.Client(project=project_id)
            bucket_name = f"{project_id}-documentai-temp"
            
            # Ensure bucket exists
            try:
                bucket = storage_client.get_bucket(bucket_name)
            except Exception:
                bucket = storage_client.create_bucket(bucket_name)

            input_blob_name = f"input/{os.path.basename(file_path)}"
            output_blob_prefix = f"output/{os.path.splitext(os.path.basename(file_path))[0]}"

            blob = bucket.blob(input_blob_name)
            blob.upload_from_filename(file_path)

            gcs_input_uri = f"gs://{bucket_name}/{input_blob_name}"
            gcs_output_uri = f"gs://{bucket_name}/{output_blob_prefix}/"

            input_config = documentai.GcsDocument(
                gcs_uri=gcs_input_uri, mime_type="application/pdf"
            )
            gcs_documents = documentai.GcsDocuments(documents=[input_config])
            input_doc = documentai.BatchDocumentsInputConfig(gcs_documents=gcs_documents)
            output_config = documentai.DocumentOutputConfig(
                gcs_output_config=documentai.DocumentOutputConfig.GcsOutputConfig(gcs_uri=gcs_output_uri)
            )

            request = documentai.BatchProcessRequest(
                name=processor_name,
                input_documents=input_doc,
                document_output_config=output_config
            )

            operation = client.batch_process_documents(request)
            operation.result(timeout=1800)  # Wait up to 30 minutes

            # Read processed result from GCS
            blobs = list(bucket.list_blobs(prefix=f"output/{os.path.splitext(os.path.basename(file_path))[0]}"))
            for blob in blobs:
                if blob.name.endswith(".json"):
                    result_json = blob.download_as_text()
                    result_obj = json.loads(result_json)
                    return result_obj.get("document", {}).get("text", "")

            return ""

    except Exception as e:
        raise Exception(f"Error processing PDF with Document AI (extended mode): {e}")


def _ocr_img(project_id: str, location: str, processor_id: str, file_path: str) -> str:
    """
    Performs OCR on an image file using Google Cloud Document AI.
    """
    try:
        opts = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=opts)
        name = client.processor_path(project_id, location, processor_id)
        
        with open(file_path, "rb") as image_file:
            image_content = image_file.read()
        
        # Determine MIME type based on file extension
        _, ext = os.path.splitext(file_path)
        mime_type_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png'
        }
        mime_type = mime_type_map.get(ext.lower(), 'image/jpeg')
        
        raw_document = documentai.RawDocument(
            content=image_content, 
            mime_type=mime_type
        )
        
        request = documentai.ProcessRequest(name=name, raw_document=raw_document)
        result = client.process_document(request=request)
        
        return result.document.text
        
    except Exception as e:
        raise Exception(f"Error processing image with Document AI: {e}")


def _get_text_from_docx(file_path: str) -> str:
    """Reads text from a local .docx file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error processing DOCX file: {e}")


def _get_text_from_msg(file_path: str) -> str:
    """Reads text from a local .msg file."""
    try:
        with extract_msg.Message(file_path) as msg:
            return f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}\n\n{msg.body}"
    except Exception as e:
        raise Exception(f"Error processing MSG file: {e}")


def _get_text_from_txt(file_path: str) -> str:
    """Reads text directly from a .txt file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise Exception(f"Error processing TXT file: {e}")


def _get_text_from_eml(file_path: str) -> str:
    """Extracts text from an .eml file."""
    try:
        with open(file_path, 'rb') as f:
            msg = email.message_from_bytes(f.read())
            main_text = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition"))
                    if "text/plain" in content_type and "attachment" not in content_disposition:
                        main_text = part.get_payload(decode=True).decode()
                        break
            else:
                main_text = msg.get_payload(decode=True).decode()
            return main_text
    except Exception as e:
        raise Exception(f"Error processing EML file: {e}")


# GCS-specific functions
# def _ocr_pdf_document_gcs(project_id: str, location: str, processor_id: str, gcs_uri: str) -> str:
#     """
#     Performs OCR on a PDF file stored in Google Cloud Storage using Document AI.
#     """
#     try:
#         client_options = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
#         client = documentai.DocumentProcessorServiceClient(client_options=client_options)
        
#         processor_name = client.processor_path(project_id, location, processor_id)
        
#         # Download the file temporarily for processing
#         bucket_name = gcs_uri.split('/')[2]
#         blob_name = '/'.join(gcs_uri.split('/')[3:])
        
#         storage_client = storage.Client(project=project_id)
#         bucket = storage_client.bucket(bucket_name)
#         blob = bucket.blob(blob_name)
        
#         # Download to temporary file
#         with tempfile.NamedTemporaryFile(suffix='.pdf') as temp_file:
#             blob.download_to_filename(temp_file.name)
            
#             with open(temp_file.name, "rb") as document_file:
#                 document_content = document_file.read()
            
#             raw_document = documentai.RawDocument(
#                 content=document_content,
#                 mime_type="application/pdf"
#             )
            
#             request = documentai.ProcessRequest(
#                 name=processor_name,
#                 raw_document=raw_document,
#                 imageless_mode=True
#             )
            
#             response = client.process_document(request=request)
#             document = response.document
            
#             return document.text
        
#     except Exception as e:
#         raise Exception(f"Error processing PDF from GCS with Document AI: {e}")

def _ocr_pdf_document_gcs(project_id: str, location: str, processor_id: str, gcs_uri: str) -> str:
    """
    Performs OCR on a PDF file stored in Google Cloud Storage using Document AI.
    Automatically switches to batch processing for large PDFs (>30 pages).
    """
    try:
        client_options = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=client_options)
        processor_name = client.processor_path(project_id, location, processor_id)

        # --- Attempt normal sync processing first ---
        try:
            input_doc = documentai.DocumentInputConfig(
                gcs_source=documentai.GcsSource(uri=gcs_uri),
                mime_type="application/pdf"
            )

            request = documentai.ProcessRequest(
                name=processor_name,
                input_documents=documentai.BatchDocumentsInputConfig(
                    gcs_documents=documentai.GcsDocuments(documents=[documentai.GcsDocument(gcs_uri=gcs_uri, mime_type="application/pdf")])
                )
            )

            response = client.process_document(request=request)
            return response.document.text

        except Exception as sync_err:
            logging.warning(f"Switching to async batch processing for GCS file: {sync_err}")

            gcs_output_uri = re.sub(r'/[^/]+$', '/output/', gcs_uri)

            input_config = documentai.GcsDocument(gcs_uri=gcs_uri, mime_type="application/pdf")
            gcs_documents = documentai.GcsDocuments(documents=[input_config])
            input_doc = documentai.BatchDocumentsInputConfig(gcs_documents=gcs_documents)
            output_config = documentai.DocumentOutputConfig(
                gcs_output_config=documentai.DocumentOutputConfig.GcsOutputConfig(gcs_uri=gcs_output_uri)
            )

            request = documentai.BatchProcessRequest(
                name=processor_name,
                input_documents=input_doc,
                document_output_config=output_config
            )

            operation = client.batch_process_documents(request)
            operation.result(timeout=1800)

            # Read result JSON from GCS
            storage_client = storage.Client(project=project_id)
            bucket_name = gcs_uri.split('/')[2]
            prefix = '/'.join(gcs_output_uri.split('/')[3:])
            bucket = storage_client.bucket(bucket_name)
            blobs = list(bucket.list_blobs(prefix=prefix))

            for blob in blobs:
                if blob.name.endswith(".json"):
                    result_json = blob.download_as_text()
                    result_obj = json.loads(result_json)
                    return result_obj.get("document", {}).get("text", "")

            return ""

    except Exception as e:
        raise Exception(f"Error processing PDF from GCS with Document AI (extended mode): {e}")


def _ocr_img_gcs(project_id: str, location: str, processor_id: str, gcs_uri: str) -> str:
    """
    Performs OCR on an image file stored in Google Cloud Storage using Document AI.
    """
    try:
        opts = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=opts)
        name = client.processor_path(project_id, location, processor_id)
        
        # Download the file temporarily for processing
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client(project=project_id)
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        # Determine MIME type based on file extension
        _, ext = os.path.splitext(blob_name)
        mime_type_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png'
        }
        mime_type = mime_type_map.get(ext.lower(), 'image/jpeg')
        
        # Download to temporary file
        with tempfile.NamedTemporaryFile(suffix=ext) as temp_file:
            blob.download_to_filename(temp_file.name)
            
            with open(temp_file.name, "rb") as image_file:
                image_content = image_file.read()
            
            raw_document = documentai.RawDocument(
                content=image_content, 
                mime_type=mime_type
            )
            
            request = documentai.ProcessRequest(name=name, raw_document=raw_document)
            result = client.process_document(request=request)
            
            return result.document.text
        
    except Exception as e:
        raise Exception(f"Error processing image from GCS with Document AI: {e}")


def _get_text_from_docx_gcs(gcs_uri: str) -> str:
    """Reads text from a .docx file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        with tempfile.NamedTemporaryFile(suffix='.docx') as temp_file:
            blob.download_to_filename(temp_file.name)
            
            doc = docx.Document(temp_file.name)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error processing DOCX file from GCS: {e}")


def _get_text_from_msg_gcs(gcs_uri: str) -> str:
    """Reads text from a .msg file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        with tempfile.NamedTemporaryFile(suffix='.msg') as temp_file:
            blob.download_to_filename(temp_file.name)
            
            with extract_msg.Message(temp_file.name) as msg:
                return f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}\n\n{msg.body}"
    except Exception as e:
        raise Exception(f"Error processing MSG file from GCS: {e}")


def _get_text_from_txt_gcs(gcs_uri: str) -> str:
    """Reads text directly from a .txt file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        return blob.download_as_text(encoding='utf-8')
    except Exception as e:
        raise Exception(f"Error processing TXT file from GCS: {e}")


def _get_text_from_eml_gcs(gcs_uri: str) -> str:
    """Extracts text from an .eml file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        with tempfile.NamedTemporaryFile(suffix='.eml') as temp_file:
            blob.download_to_filename(temp_file.name)
            
            with open(temp_file.name, 'rb') as f:
                msg = email.message_from_bytes(f.read())
                main_text = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        content_disposition = str(part.get("Content-Disposition"))
                        if "text/plain" in content_type and "attachment" not in content_disposition:
                            main_text = part.get_payload(decode=True).decode()
                            break
                else:
                    main_text = msg.get_payload(decode=True).decode()
                return main_text
    except Exception as e:
        raise Exception(f"Error processing EML file from GCS: {e}")


def _analyze_document_content(text: str, file_extension: str, filename: str) -> Dict[str, Any]:
    """
    Performs comprehensive analysis of document content including structure, language, and categorization.
    """
    try:
        if not text or not text.strip():
            return {
                "content_length": 0,
                "word_count": 0,
                "sentence_count": 0,
                "paragraph_count": 0,
                "language": "unknown",
                "document_type": "empty",
                "structure_analysis": {},
                "content_categories": [],
                "key_sections": [],
                "readability_score": 0
            }
        
        # Basic text statistics
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        content_length = len(text)
        word_count = len(words)
        sentence_count = len([s for s in sentences if s.strip()])
        paragraph_count = len(paragraphs)
        
        # Language detection (simple heuristic)
        language = _detect_language(text)
        
        # Document type classification
        document_type = _classify_document_type(text, file_extension, filename)
        
        # Structure analysis
        structure_analysis = _analyze_document_structure(text, file_extension)
        
        # Content categorization
        content_categories = _categorize_content(text, file_extension)
        
        # Key sections identification
        key_sections = _identify_key_sections(text, file_extension)
        
        # Readability analysis
        readability_score = _calculate_readability(text)
        
        return {
            "content_length": content_length,
            "word_count": word_count,
            "sentence_count": sentence_count,
            "paragraph_count": paragraph_count,
            "language": language,
            "document_type": document_type,
            "structure_analysis": structure_analysis,
            "content_categories": content_categories,
            "key_sections": key_sections,
            "readability_score": readability_score,
            "average_words_per_sentence": word_count / max(sentence_count, 1),
            "average_sentences_per_paragraph": sentence_count / max(paragraph_count, 1)
        }
        
    except Exception as e:
        logging.error(f"Error in document content analysis: {str(e)}")
        return {
            "content_length": len(text) if text else 0,
            "word_count": 0,
            "sentence_count": 0,
            "paragraph_count": 0,
            "language": "unknown",
            "document_type": "unknown",
            "structure_analysis": {},
            "content_categories": [],
            "key_sections": [],
            "readability_score": 0,
            "analysis_error": str(e)
        }


def _generate_file_metadata(filename: str, file_extension: str, file_size: int, 
                          processing_method: str, processing_time: float, 
                          is_gcs_url: bool, file_path: str) -> Dict[str, Any]:
    """
    Generates comprehensive file metadata including processing information.
    """
    try:
        # Generate file hash for uniqueness
        file_hash = hashlib.md5(f"{filename}_{file_size}_{datetime.datetime.now().isoformat()}".encode()).hexdigest()[:16]
        
        # Determine file category
        file_category = _get_file_category(file_extension)
        
        # Processing confidence based on method
        confidence_scores = {
            "Google Cloud Document AI": 0.95,
            "python-docx": 0.98,
            "extract-msg": 0.90,
            "Python email library": 0.85,
            "Direct text reading": 0.99
        }
        
        confidence = confidence_scores.get(processing_method.split('(')[0].strip(), 0.80)
        
        return {
            "filename": filename,
            "file_extension": file_extension,
            "file_category": file_category,
            "file_size_bytes": file_size,
            "file_size_mb": round(file_size / (1024 * 1024), 2),
            "processing_method": processing_method,
            "processing_time_seconds": round(processing_time, 3),
            "processing_confidence": confidence,
            "is_gcs_url": is_gcs_url,
            "file_path": file_path,
            "file_hash": file_hash,
            "timestamp": datetime.datetime.now().isoformat(),
            "supported_operations": _get_supported_operations(file_extension)
        }
        
    except Exception as e:
        logging.error(f"Error generating file metadata: {str(e)}")
        return {
            "filename": filename,
            "file_extension": file_extension,
            "file_size_bytes": file_size,
            "processing_method": processing_method,
            "processing_time_seconds": processing_time,
            "is_gcs_url": is_gcs_url,
            "file_path": file_path,
            "timestamp": datetime.datetime.now().isoformat(),
            "metadata_error": str(e)
        }


def _analyze_content_quality(text: str, file_extension: str) -> Dict[str, Any]:
    """
    Analyzes the quality of extracted content including OCR confidence and text integrity.
    """
    try:
        if not text or not text.strip():
            return {
                "text_quality_score": 0,
                "ocr_confidence": 0,
                "encoding_issues": 0,
                "special_characters_ratio": 0,
                "whitespace_ratio": 0,
                "line_breaks_ratio": 0,
                "quality_issues": ["Empty content"],
                "text_integrity": "poor"
            }
        
        # Calculate various quality metrics
        total_chars = len(text)
        whitespace_chars = text.count(' ') + text.count('\t') + text.count('\n')
        line_breaks = text.count('\n')
        special_chars = len(re.findall(r'[^\w\s]', text))
        
        # Check for common OCR issues
        ocr_issues = []
        if re.search(r'[a-zA-Z]{1,2}\s+[a-zA-Z]{1,2}', text):  # Split words
            ocr_issues.append("Possible word splitting")
        if re.search(r'[0-9]+\s+[0-9]+', text):  # Split numbers
            ocr_issues.append("Possible number splitting")
        if re.search(r'[^\x00-\x7F]', text):  # Non-ASCII characters
            ocr_issues.append("Non-ASCII characters detected")
        
        # Calculate quality score (0-100)
        whitespace_ratio = whitespace_chars / total_chars if total_chars > 0 else 0
        special_char_ratio = special_chars / total_chars if total_chars > 0 else 0
        line_break_ratio = line_breaks / total_chars if total_chars > 0 else 0
        
        # Quality scoring
        quality_score = 100
        if whitespace_ratio > 0.3:
            quality_score -= 20
        if special_char_ratio > 0.1:
            quality_score -= 15
        if line_break_ratio > 0.1:
            quality_score -= 10
        if len(ocr_issues) > 0:
            quality_score -= len(ocr_issues) * 5
        
        quality_score = max(0, min(100, quality_score))
        
        # Determine text integrity
        if quality_score >= 90:
            integrity = "excellent"
        elif quality_score >= 75:
            integrity = "good"
        elif quality_score >= 60:
            integrity = "fair"
        else:
            integrity = "poor"
        
        return {
            "text_quality_score": round(quality_score, 2),
            "ocr_confidence": round(quality_score / 100, 3),
            "encoding_issues": 1 if re.search(r'[^\x00-\x7F]', text) else 0,
            "special_characters_ratio": round(special_char_ratio, 3),
            "whitespace_ratio": round(whitespace_ratio, 3),
            "line_breaks_ratio": round(line_break_ratio, 3),
            "quality_issues": ocr_issues,
            "text_integrity": integrity,
            "total_characters": total_chars,
            "readable_characters": total_chars - special_chars
        }
        
    except Exception as e:
        logging.error(f"Error analyzing content quality: {str(e)}")
        return {
            "text_quality_score": 0,
            "ocr_confidence": 0,
            "encoding_issues": 1,
            "special_characters_ratio": 0,
            "whitespace_ratio": 0,
            "line_breaks_ratio": 0,
            "quality_issues": ["Analysis error"],
            "text_integrity": "unknown",
            "analysis_error": str(e)
        }


def _extract_key_information(text: str, file_extension: str, filename: str) -> Dict[str, Any]:
    """
    Extracts key information from the document based on its type and content.
    """
    try:
        if not text or not text.strip():
            return {
                "document_title": "",
                "key_entities": [],
                "important_numbers": [],
                "dates": [],
                "email_addresses": [],
                "urls": [],
                "phone_numbers": [],
                "financial_data": {},
                "contact_information": {},
                "business_terms": [],
                "summary": "No content to analyze"
            }
        
        # Extract various types of information
        email_addresses = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', text)
        phone_numbers = re.findall(r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})', text)
        dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b', text)
        
        # Extract financial data
        financial_data = _extract_financial_data(text)
        
        # Extract business terms
        business_terms = _extract_business_terms(text)
        
        # Extract key entities (names, companies, etc.)
        key_entities = _extract_key_entities(text)
        
        # Extract important numbers
        important_numbers = _extract_important_numbers(text)
        
        # Generate document title
        document_title = _generate_document_title(text, filename)
        
        # Extract contact information
        contact_info = _extract_contact_information(text, email_addresses, phone_numbers)
        
        # Generate summary
        summary = _generate_document_summary(text, file_extension)
        
        return {
            "document_title": document_title,
            "key_entities": key_entities,
            "important_numbers": important_numbers,
            "dates": dates,
            "email_addresses": list(set(email_addresses)),
            "urls": list(set(urls)),
            "phone_numbers": [''.join(p) for p in phone_numbers],
            "financial_data": financial_data,
            "contact_information": contact_info,
            "business_terms": business_terms,
            "summary": summary,
            "extraction_confidence": _calculate_extraction_confidence(text, file_extension)
        }
        
    except Exception as e:
        logging.error(f"Error extracting key information: {str(e)}")
        return {
            "document_title": "",
            "key_entities": [],
            "important_numbers": [],
            "dates": [],
            "email_addresses": [],
            "urls": [],
            "phone_numbers": [],
            "financial_data": {},
            "contact_information": {},
            "business_terms": [],
            "summary": "Error in information extraction",
            "extraction_error": str(e)
        }


# Helper functions for analysis
def _detect_language(text: str) -> str:
    """Simple language detection based on common words."""
    english_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
    spanish_words = ['el', 'la', 'de', 'que', 'y', 'a', 'en', 'un', 'es', 'se', 'no', 'te', 'lo', 'le']
    french_words = ['le', 'la', 'de', 'et', 'à', 'un', 'il', 'être', 'et', 'en', 'avoir', 'que', 'pour']
    
    text_lower = text.lower()
    english_count = sum(1 for word in english_words if word in text_lower)
    spanish_count = sum(1 for word in spanish_words if word in text_lower)
    french_count = sum(1 for word in french_words if word in text_lower)
    
    if english_count > spanish_count and english_count > french_count:
        return "english"
    elif spanish_count > french_count:
        return "spanish"
    elif french_count > 0:
        return "french"
    else:
        return "unknown"


def _classify_document_type(text: str, file_extension: str, filename: str) -> str:
    """Classify the type of document based on content and filename."""
    text_lower = text.lower()
    filename_lower = filename.lower()
    
    # Check for specific document types
    if any(word in text_lower for word in ['pitch deck', 'pitch deck', 'investor presentation', 'funding deck']):
        return "pitch_deck"
    elif any(word in text_lower for word in ['business plan', 'business proposal', 'executive summary']):
        return "business_plan"
    elif any(word in text_lower for word in ['financial statement', 'income statement', 'balance sheet', 'cash flow']):
        return "financial_document"
    elif any(word in text_lower for word in ['contract', 'agreement', 'terms and conditions', 'legal document']):
        return "legal_document"
    elif any(word in text_lower for word in ['resume', 'cv', 'curriculum vitae', 'personal profile']):
        return "resume"
    elif any(word in text_lower for word in ['email', 'message', 'correspondence']):
        return "email"
    elif any(word in filename_lower for word in ['invoice', 'receipt', 'bill']):
        return "invoice"
    elif any(word in filename_lower for word in ['report', 'analysis', 'study']):
        return "report"
    else:
        return "general_document"


def _analyze_document_structure(text: str, file_extension: str) -> Dict[str, Any]:
    """Analyze the structural elements of the document."""
    try:
        # Find headers (lines that are all caps or start with numbers)
        headers = re.findall(r'^[A-Z\s]{3,}$|^\d+\.?\s+[A-Z]', text, re.MULTILINE)
        
        # Find bullet points
        bullet_points = re.findall(r'^[\s]*[•\-\*]\s+', text, re.MULTILINE)
        
        # Find numbered lists
        numbered_lists = re.findall(r'^\d+\.\s+', text, re.MULTILINE)
        
        # Find tables (lines with multiple spaces or tabs)
        table_lines = re.findall(r'^.*\s{3,}.*$', text, re.MULTILINE)
        
        # Find potential sections
        sections = re.findall(r'^[A-Z][A-Za-z\s]{5,}:?$', text, re.MULTILINE)
        
        return {
            "headers_count": len(headers),
            "bullet_points_count": len(bullet_points),
            "numbered_lists_count": len(numbered_lists),
            "table_lines_count": len(table_lines),
            "sections_count": len(sections),
            "has_structure": len(headers) > 0 or len(bullet_points) > 0 or len(numbered_lists) > 0,
            "structure_type": _determine_structure_type(headers, bullet_points, numbered_lists)
        }
        
    except Exception as e:
        return {"error": str(e)}


def _categorize_content(text: str, file_extension: str) -> List[str]:
    """Categorize the content based on keywords and patterns."""
    categories = []
    text_lower = text.lower()
    
    # Business categories
    if any(word in text_lower for word in ['revenue', 'profit', 'sales', 'income', 'financial']):
        categories.append("financial")
    if any(word in text_lower for word in ['market', 'customer', 'user', 'target', 'demographic']):
        categories.append("market_analysis")
    if any(word in text_lower for word in ['product', 'service', 'feature', 'development']):
        categories.append("product")
    if any(word in text_lower for word in ['team', 'employee', 'staff', 'founder', 'ceo']):
        categories.append("team")
    if any(word in text_lower for word in ['funding', 'investment', 'investor', 'capital', 'raise']):
        categories.append("funding")
    if any(word in text_lower for word in ['technology', 'tech', 'software', 'platform', 'app']):
        categories.append("technology")
    if any(word in text_lower for word in ['legal', 'contract', 'agreement', 'terms', 'compliance']):
        categories.append("legal")
    
    return categories if categories else ["general"]


def _identify_key_sections(text: str, file_extension: str) -> List[Dict[str, str]]:
    """Identify key sections in the document."""
    sections = []
    
    # Common section headers
    section_patterns = [
        r'(executive\s+summary)',
        r'(introduction)',
        r'(problem\s+statement)',
        r'(solution)',
        r'(market\s+analysis)',
        r'(business\s+model)',
        r'(financial\s+projections?)',
        r'(team)',
        r'(funding\s+requirements?)',
        r'(conclusion)',
        r'(appendix)',
        r'(references?)'
    ]
    
    for pattern in section_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            sections.append({
                "title": match.group(1).strip(),
                "position": match.start(),
                "type": "section_header"
            })
    
    return sections


def _calculate_readability(text: str) -> float:
    """Calculate a simple readability score."""
    try:
        sentences = re.split(r'[.!?]+', text)
        words = text.split()
        
        if len(sentences) == 0 or len(words) == 0:
            return 0
        
        avg_sentence_length = len(words) / len(sentences)
        avg_word_length = sum(len(word) for word in words) / len(words)
        
        # Simple readability formula (lower is more readable)
        readability = (avg_sentence_length * 0.5) + (avg_word_length * 0.3)
        return round(readability, 2)
        
    except:
        return 0


def _get_file_category(file_extension: str) -> str:
    """Get the category of the file based on its extension."""
    categories = {
        '.pdf': 'document',
        '.docx': 'document',
        '.txt': 'text',
        '.png': 'image',
        '.jpg': 'image',
        '.jpeg': 'image',
        '.tiff': 'image',
        '.eml': 'email',
        '.msg': 'email'
    }
    return categories.get(file_extension, 'unknown')


def _get_supported_operations(file_extension: str) -> List[str]:
    """Get list of supported operations for the file type."""
    operations = {
        '.pdf': ['ocr', 'text_extraction', 'structure_analysis'],
        '.docx': ['text_extraction', 'structure_analysis', 'metadata_extraction'],
        '.txt': ['text_extraction', 'content_analysis'],
        '.png': ['ocr', 'text_extraction'],
        '.jpg': ['ocr', 'text_extraction'],
        '.jpeg': ['ocr', 'text_extraction'],
        '.tiff': ['ocr', 'text_extraction'],
        '.eml': ['email_parsing', 'text_extraction', 'header_extraction'],
        '.msg': ['email_parsing', 'text_extraction', 'header_extraction']
    }
    return operations.get(file_extension, ['text_extraction'])


def _extract_financial_data(text: str) -> Dict[str, Any]:
    """Extract financial data from the text."""
    financial_data = {
        "currency_amounts": [],
        "percentages": [],
        "financial_terms": []
    }
    
    # Extract currency amounts
    currency_patterns = [
        r'\$[\d,]+(?:\.\d{2})?',
        r'USD\s*[\d,]+(?:\.\d{2})?',
        r'EUR\s*[\d,]+(?:\.\d{2})?',
        r'GBP\s*[\d,]+(?:\.\d{2})?'
    ]
    
    for pattern in currency_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        financial_data["currency_amounts"].extend(matches)
    
    # Extract percentages
    percentage_matches = re.findall(r'\d+(?:\.\d+)?%', text)
    financial_data["percentages"] = percentage_matches
    
    # Extract financial terms
    financial_terms = [
        'revenue', 'profit', 'loss', 'income', 'expense', 'cost', 'price',
        'valuation', 'investment', 'funding', 'capital', 'equity', 'debt',
        'cash flow', 'margin', 'growth', 'ROI', 'EBITDA'
    ]
    
    found_terms = []
    for term in financial_terms:
        if re.search(r'\b' + re.escape(term) + r'\b', text, re.IGNORECASE):
            found_terms.append(term)
    
    financial_data["financial_terms"] = found_terms
    
    return financial_data


def _extract_business_terms(text: str) -> List[str]:
    """Extract business-related terms from the text."""
    business_terms = [
        'startup', 'company', 'corporation', 'LLC', 'Inc', 'Ltd',
        'CEO', 'CTO', 'CFO', 'founder', 'co-founder', 'president',
        'venture capital', 'angel investor', 'seed funding', 'Series A',
        'IPO', 'acquisition', 'merger', 'partnership', 'collaboration',
        'market', 'customer', 'user', 'client', 'revenue model',
        'business model', 'value proposition', 'competitive advantage'
    ]
    
    found_terms = []
    for term in business_terms:
        if re.search(r'\b' + re.escape(term) + r'\b', text, re.IGNORECASE):
            found_terms.append(term)
    
    return found_terms


def _extract_key_entities(text: str) -> List[str]:
    """Extract key entities like names, companies, etc."""
    entities = []
    
    # Extract potential company names (words that start with capital letters)
    company_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b'
    potential_companies = re.findall(company_pattern, text)
    
    # Filter out common words
    common_words = {'The', 'This', 'That', 'There', 'Here', 'Where', 'When', 'How', 'What', 'Why'}
    companies = [comp for comp in potential_companies if comp not in common_words and len(comp) > 3]
    
    entities.extend(companies[:10])  # Limit to top 10
    
    return entities


def _extract_important_numbers(text: str) -> List[str]:
    """Extract important numbers from the text."""
    numbers = []
    
    # Extract large numbers (likely important)
    large_numbers = re.findall(r'\b\d{4,}\b', text)
    numbers.extend(large_numbers)
    
    # Extract decimal numbers
    decimal_numbers = re.findall(r'\b\d+\.\d+\b', text)
    numbers.extend(decimal_numbers)
    
    return numbers[:20]  # Limit to top 20


def _generate_document_title(text: str, filename: str) -> str:
    """Generate a document title based on content and filename."""
    # Try to find a title in the first few lines
    lines = text.split('\n')[:10]
    for line in lines:
        line = line.strip()
        if len(line) > 10 and len(line) < 100 and not line.isupper():
            return line
    
    # Fall back to filename without extension
    return os.path.splitext(filename)[0]


def _extract_contact_information(text: str, emails: List[str], phones: List[str]) -> Dict[str, Any]:
    """Extract contact information from the text."""
    contact_info = {
        "emails": emails,
        "phone_numbers": phones,
        "addresses": [],
        "websites": []
    }
    
    # Extract addresses (simple pattern)
    address_pattern = r'\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd)'
    addresses = re.findall(address_pattern, text, re.IGNORECASE)
    contact_info["addresses"] = addresses
    
    # Extract websites
    website_pattern = r'(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:/[^\s]*)?'
    websites = re.findall(website_pattern, text)
    contact_info["websites"] = websites
    
    return contact_info


def _generate_document_summary(text: str, file_extension: str) -> str:
    """Generate a summary of the document."""
    if not text or len(text) < 100:
        return "Document contains minimal content."
    
    # Take first few sentences as summary
    sentences = re.split(r'[.!?]+', text)
    summary_sentences = sentences[:3]
    summary = '. '.join(sentence.strip() for sentence in summary_sentences if sentence.strip())
    
    if len(summary) > 200:
        summary = summary[:200] + "..."
    
    return summary


def _calculate_extraction_confidence(text: str, file_extension: str) -> float:
    """Calculate confidence in the extraction quality."""
    if not text:
        return 0.0
    
    confidence = 0.8  # Base confidence
    
    # Adjust based on file type
    if file_extension in ['.txt', '.docx']:
        confidence += 0.1
    elif file_extension in ['.pdf', '.tiff']:
        confidence += 0.05
    elif file_extension in ['.png', '.jpg', '.jpeg']:
        confidence += 0.0
    elif file_extension in ['.eml', '.msg']:
        confidence += 0.05
    
    # Adjust based on text quality
    if len(text) > 1000:
        confidence += 0.05
    if re.search(r'[a-zA-Z]', text):  # Contains letters
        confidence += 0.05
    
    return min(1.0, confidence)


def _determine_structure_type(headers: List[str], bullets: List[str], numbers: List[str]) -> str:
    """Determine the type of document structure."""
    if len(headers) > 5:
        return "formal_document"
    elif len(bullets) > 5:
        return "bullet_point_format"
    elif len(numbers) > 5:
        return "numbered_list"
    elif len(headers) > 0:
        return "semi_structured"
    else:
        return "unstructured"
