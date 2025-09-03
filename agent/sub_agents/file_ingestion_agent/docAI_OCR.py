# google-api-core               
# google-auth                   
# google-cloud-aiplatform       
# google-cloud-bigquery         
# google-cloud-core             
# google-cloud-documentai       
# google-cloud-resource-manager 
# google-cloud-storage          
# google-cloud-vision           
# google-crc32c                 
# google-genai                  
# google-resumable-media        
# googleapis-common-protos  

# The above might be required - just adding for reference

import json
import re
import os
from google.cloud import vision
from google.cloud import storage
import docx
import extract_msg
from google.api_core.client_options import ClientOptions
from google.cloud import documentai_v1beta3 as documentai
import os
from docx import Document
import email
import mimetypes



def ocr_pdf_document(project_id: str, location: str, processor_id: str, file_path: str):
    """
    Performs OCR on a local PDF file using the Google Cloud Document AI API.

    Args:
        project_id: The Google Cloud project ID.
        location: The location of the processor (e.g., "us" or "eu").
        processor_id: The ID of the processor to use.
        file_path: The full path to the PDF file to be processed.
    """
    try:
        # Create a client with the correct regional endpoint.
        # This resolves the `AttributeError` by importing ClientOptions from the correct
        # google.api_core module and using it to set the API endpoint.
        client_options = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=client_options)

        # The full resource name of the processor
        processor_name = client.processor_path(project_id, location, processor_id)

        # Read the file into memory.
        with open(file_path, "rb") as document_file:
            document_content = document_file.read()
        
        # Create a RawDocument object. The mime_type is crucial for the API.
        raw_document = documentai.RawDocument(
            content=document_content,
            mime_type="application/pdf"
        )
        
        # Configure the process request.
        request = documentai.ProcessRequest(
            name=processor_name,
            raw_document=raw_document,
            imageless_mode=True
        )

        # Execute the request and get the response.
        print(f"Sending request for OCR on '{os.path.basename(file_path)}'...")
        response = client.process_document(request=request)
        print("Processing complete.")

        # The Document object is a representation of the document,
        # with all the extracted information.
        document = response.document
        
        # Print the extracted text.
        #print("\n--- Extracted Text ---")
        return document.text

        # Optionally, you can also access other extracted information like entities or pages
        # for a more structured output. This is a basic example for text extraction.

    except Exception as e:
        print(f"An error occurred: {e}")

def ocr_img(project_id, location, processor_id, file_path):
    """
    Performs OCR on a file using Google Cloud Document AI.
    """
    opts = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
    client = documentai.DocumentProcessorServiceClient(client_options=opts)
    name = client.processor_path(project_id, location, processor_id)
    
    with open(file_path, "rb") as image_file:
        image_content = image_file.read()
    
    raw_document = documentai.RawDocument(
        content=image_content, mime_type="image/jpeg" # or other mime type
    )
    
    request = documentai.ProcessRequest(name=name, raw_document=raw_document)
    result = client.process_document(request=request)
    
    # The output is a complex Document object, not just raw text
    return result.document.text


def get_text_from_docx(file_path):
    """Reads text from a local .docx file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error processing DOCX file: {e}"

def get_text_from_msg(file_path):
    """Reads text from a local .msg file."""
    try:
        with extract_msg.Message(file_path) as msg:
            return f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}\n\n{msg.body}"
    except Exception as e:
        return f"Error processing MSG file: {e}"

def get_text_from_txt(file_path):
    """Reads text directly from a .txt file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Error processing TXT file: {e}"

def get_text_from_eml(file_path):
    """Extracts text from an .eml file."""
    try:
        with open(file_path, 'rb') as f:
            msg = email.message_from_bytes(f.read())
            main_text = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition"))
                    if "text/plain" in content_type and "attachment" not in content_disposition:
                        main_text = part.get_payload(decode=True).decode()
                        break
            else:
                main_text = msg.get_payload(decode=True).decode()
            return main_text
    except Exception as e:
        return f"Error processing EML file: {e}"

def ingest_document(file_path):
    """
    Intelligently ingests a document by detecting its format and
    using the appropriate method to extract text.
    """
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == '.docx':
        print(f"Ingesting {file_extension} document using python-docx.")
        text = get_text_from_docx(file_path)
    elif file_extension == '.eml':
        print(f"Ingesting {file_extension} document using py in-build email library.")
        text = get_text_from_eml(file_path)
    elif file_extension == '.msg':
        print(f"Ingesting {file_extension} document using extract-msg.")
        text = get_text_from_msg(file_path)
    elif file_extension == '.txt':
        print(f"Ingesting {file_extension} document directly.")
        text = get_text_from_txt(file_path)
    else:
        print(f"Unsupported file format: {file_extension}")
        text = ""
    return text

if __name__ == '__main__':
    # Define the path to your resources directory
    resources_path = '/Users/kirupa/Documents/projects/agenticai/genAIHackTrial/resources_subAgentA'
    
    # Define the GCS bucket and folders
    #gcs_bucket = "genai-hackathon-2025"
    #gcs_upload_folder = 'temp_uploads/'
    #gcs_output_folder = 'vision-output/'
    
    # Loop through all files in the resources directory
    for filename in os.listdir(resources_path):
        local_file_path = os.path.join(resources_path, filename)
        
        # Skip directories and hidden files
        if os.path.isdir(local_file_path) or filename.startswith('.'):
            continue
            
        print(f"\n--- Processing file: {filename} ---")
        
        _, file_extension = os.path.splitext(filename)
        file_extension = file_extension.lower()

        extracted_text = ""
        
        PROJECT_ID = "quiet-sum-470418-r7"
        LOCATION = "us"
        PROCESSOR_ID = "abb1ab40cbac8a9c"
        FILE_PATH = f"/Users/kirupa/Documents/projects/agenticai/genAIHackTrial/resources_subAgentA/{filename}" # e.g., "NVCPitchDeckTemplate.pdf"

        if file_extension in ['.pdf', '.tiff']:
            extracted_text = ocr_pdf_document(PROJECT_ID, LOCATION, PROCESSOR_ID, FILE_PATH)

        elif file_extension in ['.png', '.jpeg', '.jpg']:
            # For images, use the Vision API directly from the local file
            # print(f"Ingesting {file_extension} file using Vision API.")
            # client = vision.ImageAnnotatorClient()
            # with open(local_file_path, "rb") as image_file:
            #     content = image_file.read()
            # image = vision.Image(content=content)
            # response = client.text_detection(image=image)
            # texts = response.text_annotations
            # if texts:
            #     extracted_text = texts[0].description
            # else:
            #     extracted_text = "No text detected in the image."
            extracted_text = ocr_img(PROJECT_ID, LOCATION, PROCESSOR_ID, FILE_PATH)
        
        else:
            # For all other formats, use the ingest_document function
            extracted_text = ingest_document(local_file_path)

        # Print the extracted text (or a snippet)
        print("\n--- Extracted Text ---")
        print(extracted_text)
