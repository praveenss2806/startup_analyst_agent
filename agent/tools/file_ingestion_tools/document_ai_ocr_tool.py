import json
import re
import os
from google.cloud import vision
from google.cloud import storage
import docx
import extract_msg
from google.api_core.client_options import ClientOptions
from google.cloud import documentai_v1beta3 as documentai
from docx import Document
import email
import mimetypes
import tempfile
from typing import Dict, Any
from google.adk.tools import ToolContext


def document_ai_ocr_tool(file_path: str, tool_context: ToolContext) -> Dict[str, Any]:
    """
    Intelligently processes documents using Google Cloud Document AI OCR and other text extraction methods.
    Supports PDF, images, DOCX, MSG, EML, and TXT files from local paths or GCS URLs.

    Args:
        file_path (str): Path to the input file (local path or GCS URL).
        tool_context (ToolContext): The tool context containing state information.
        
    Returns:
        Dict[str, Any]: A dictionary containing:
            - status: Success or Failure
            - extracted_text: The extracted text content
            - file_type: The type of file processed
    """
    try:
        # Google Cloud Document AI configuration
        PROJECT_ID = "quiet-sum-470418-r7"
        LOCATION = "us"
        PROCESSOR_ID = "abb1ab40cbac8a9c"
        
        # Check if file_path is a GCS URL or local path
        is_gcs_url = file_path.startswith('gs://') or file_path.startswith('https://storage.googleapis.com/')
        
        if is_gcs_url:
            # Extract file extension from GCS URL
            if file_path.startswith('https://storage.googleapis.com/'):
                # Convert public URL to GCS URI format
                gcs_uri = file_path.replace('https://storage.googleapis.com/', 'gs://')
            else:
                gcs_uri = file_path
            
            # Extract filename from GCS URI
            filename = gcs_uri.split('/')[-1]
            _, file_extension = os.path.splitext(filename)
        else:
            # Local file path
            _, file_extension = os.path.splitext(file_path)
            
        file_extension = file_extension.lower()
        
        extracted_text = ""
        processing_method = ""
        
        if file_extension in ['.pdf', '.tiff']:
            if is_gcs_url:
                extracted_text = _ocr_pdf_document_gcs(PROJECT_ID, LOCATION, PROCESSOR_ID, gcs_uri)
                processing_method = "Google Cloud Document AI (PDF/TIFF from GCS)"
            else:
                extracted_text = _ocr_pdf_document(PROJECT_ID, LOCATION, PROCESSOR_ID, file_path)
                processing_method = "Google Cloud Document AI (PDF/TIFF)"
            
        elif file_extension in ['.png', '.jpeg', '.jpg']:
            if is_gcs_url:
                extracted_text = _ocr_img_gcs(PROJECT_ID, LOCATION, PROCESSOR_ID, gcs_uri)
                processing_method = "Google Cloud Document AI (Image from GCS)"
            else:
                extracted_text = _ocr_img(PROJECT_ID, LOCATION, PROCESSOR_ID, file_path)
                processing_method = "Google Cloud Document AI (Image)"
            
        elif file_extension == '.docx':
            if is_gcs_url:
                extracted_text = _get_text_from_docx_gcs(gcs_uri)
                processing_method = "python-docx (from GCS)"
            else:
                extracted_text = _get_text_from_docx(file_path)
                processing_method = "python-docx"
            
        elif file_extension == '.eml':
            if is_gcs_url:
                extracted_text = _get_text_from_eml_gcs(gcs_uri)
                processing_method = "Python email library (from GCS)"
            else:
                extracted_text = _get_text_from_eml(file_path)
                processing_method = "Python email library"
            
        elif file_extension == '.msg':
            if is_gcs_url:
                extracted_text = _get_text_from_msg_gcs(gcs_uri)
                processing_method = "extract-msg (from GCS)"
            else:
                extracted_text = _get_text_from_msg(file_path)
                processing_method = "extract-msg"
            
        elif file_extension == '.txt':
            if is_gcs_url:
                extracted_text = _get_text_from_txt_gcs(gcs_uri)
                processing_method = "Direct text reading (from GCS)"
            else:
                extracted_text = _get_text_from_txt(file_path)
                processing_method = "Direct text reading"
            
        else:
            return {
                "status": "failure",
                "error": f"Unsupported file format: {file_extension}"
            }
        
        # Store extracted data in tool context
        tool_context.state['startup_information'] = extracted_text
        tool_context.state['file_path'] = file_path
        tool_context.state['file_type'] = file_extension
        tool_context.state['processing_method'] = processing_method
        
        return {
            "status": "success",
            "extracted_text": extracted_text,
            "file_type": file_extension,
            "processing_method": processing_method
        }
        
    except Exception as e:
        return {
            "status": "failure", 
            "error": str(e)
        }


def _ocr_pdf_document(project_id: str, location: str, processor_id: str, file_path: str) -> str:
    """
    Performs OCR on a local PDF file using the Google Cloud Document AI API.
    """
    try:
        client_options = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=client_options)
        
        processor_name = client.processor_path(project_id, location, processor_id)
        
        with open(file_path, "rb") as document_file:
            document_content = document_file.read()
        
        raw_document = documentai.RawDocument(
            content=document_content,
            mime_type="application/pdf"
        )
        
        request = documentai.ProcessRequest(
            name=processor_name,
            raw_document=raw_document,
            imageless_mode=True
        )
        
        response = client.process_document(request=request)
        document = response.document
        
        return document.text
        
    except Exception as e:
        raise Exception(f"Error processing PDF with Document AI: {e}")


def _ocr_img(project_id: str, location: str, processor_id: str, file_path: str) -> str:
    """
    Performs OCR on an image file using Google Cloud Document AI.
    """
    try:
        opts = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=opts)
        name = client.processor_path(project_id, location, processor_id)
        
        with open(file_path, "rb") as image_file:
            image_content = image_file.read()
        
        # Determine MIME type based on file extension
        _, ext = os.path.splitext(file_path)
        mime_type_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png'
        }
        mime_type = mime_type_map.get(ext.lower(), 'image/jpeg')
        
        raw_document = documentai.RawDocument(
            content=image_content, 
            mime_type=mime_type
        )
        
        request = documentai.ProcessRequest(name=name, raw_document=raw_document)
        result = client.process_document(request=request)
        
        return result.document.text
        
    except Exception as e:
        raise Exception(f"Error processing image with Document AI: {e}")


def _get_text_from_docx(file_path: str) -> str:
    """Reads text from a local .docx file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error processing DOCX file: {e}")


def _get_text_from_msg(file_path: str) -> str:
    """Reads text from a local .msg file."""
    try:
        with extract_msg.Message(file_path) as msg:
            return f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}\n\n{msg.body}"
    except Exception as e:
        raise Exception(f"Error processing MSG file: {e}")


def _get_text_from_txt(file_path: str) -> str:
    """Reads text directly from a .txt file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise Exception(f"Error processing TXT file: {e}")


def _get_text_from_eml(file_path: str) -> str:
    """Extracts text from an .eml file."""
    try:
        with open(file_path, 'rb') as f:
            msg = email.message_from_bytes(f.read())
            main_text = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition"))
                    if "text/plain" in content_type and "attachment" not in content_disposition:
                        main_text = part.get_payload(decode=True).decode()
                        break
            else:
                main_text = msg.get_payload(decode=True).decode()
            return main_text
    except Exception as e:
        raise Exception(f"Error processing EML file: {e}")


# GCS-specific functions
def _ocr_pdf_document_gcs(project_id: str, location: str, processor_id: str, gcs_uri: str) -> str:
    """
    Performs OCR on a PDF file stored in Google Cloud Storage using Document AI.
    """
    try:
        client_options = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=client_options)
        
        processor_name = client.processor_path(project_id, location, processor_id)
        
        # Download the file temporarily for processing
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client(project=project_id)
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        # Download to temporary file
        with tempfile.NamedTemporaryFile(suffix='.pdf') as temp_file:
            blob.download_to_filename(temp_file.name)
            
            with open(temp_file.name, "rb") as document_file:
                document_content = document_file.read()
            
            raw_document = documentai.RawDocument(
                content=document_content,
                mime_type="application/pdf"
            )
            
            request = documentai.ProcessRequest(
                name=processor_name,
                raw_document=raw_document,
                imageless_mode=True
            )
            
            response = client.process_document(request=request)
            document = response.document
            
            return document.text
        
    except Exception as e:
        raise Exception(f"Error processing PDF from GCS with Document AI: {e}")


def _ocr_img_gcs(project_id: str, location: str, processor_id: str, gcs_uri: str) -> str:
    """
    Performs OCR on an image file stored in Google Cloud Storage using Document AI.
    """
    try:
        opts = ClientOptions(api_endpoint=f"{location}-documentai.googleapis.com")
        client = documentai.DocumentProcessorServiceClient(client_options=opts)
        name = client.processor_path(project_id, location, processor_id)
        
        # Download the file temporarily for processing
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client(project=project_id)
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        # Determine MIME type based on file extension
        _, ext = os.path.splitext(blob_name)
        mime_type_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png'
        }
        mime_type = mime_type_map.get(ext.lower(), 'image/jpeg')
        
        # Download to temporary file
        with tempfile.NamedTemporaryFile(suffix=ext) as temp_file:
            blob.download_to_filename(temp_file.name)
            
            with open(temp_file.name, "rb") as image_file:
                image_content = image_file.read()
            
            raw_document = documentai.RawDocument(
                content=image_content, 
                mime_type=mime_type
            )
            
            request = documentai.ProcessRequest(name=name, raw_document=raw_document)
            result = client.process_document(request=request)
            
            return result.document.text
        
    except Exception as e:
        raise Exception(f"Error processing image from GCS with Document AI: {e}")


def _get_text_from_docx_gcs(gcs_uri: str) -> str:
    """Reads text from a .docx file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        with tempfile.NamedTemporaryFile(suffix='.docx') as temp_file:
            blob.download_to_filename(temp_file.name)
            
            doc = docx.Document(temp_file.name)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error processing DOCX file from GCS: {e}")


def _get_text_from_msg_gcs(gcs_uri: str) -> str:
    """Reads text from a .msg file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        with tempfile.NamedTemporaryFile(suffix='.msg') as temp_file:
            blob.download_to_filename(temp_file.name)
            
            with extract_msg.Message(temp_file.name) as msg:
                return f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}\n\n{msg.body}"
    except Exception as e:
        raise Exception(f"Error processing MSG file from GCS: {e}")


def _get_text_from_txt_gcs(gcs_uri: str) -> str:
    """Reads text directly from a .txt file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        return blob.download_as_text(encoding='utf-8')
    except Exception as e:
        raise Exception(f"Error processing TXT file from GCS: {e}")


def _get_text_from_eml_gcs(gcs_uri: str) -> str:
    """Extracts text from an .eml file stored in Google Cloud Storage."""
    try:
        bucket_name = gcs_uri.split('/')[2]
        blob_name = '/'.join(gcs_uri.split('/')[3:])
        
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        with tempfile.NamedTemporaryFile(suffix='.eml') as temp_file:
            blob.download_to_filename(temp_file.name)
            
            with open(temp_file.name, 'rb') as f:
                msg = email.message_from_bytes(f.read())
                main_text = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        content_disposition = str(part.get("Content-Disposition"))
                        if "text/plain" in content_type and "attachment" not in content_disposition:
                            main_text = part.get_payload(decode=True).decode()
                            break
                else:
                    main_text = msg.get_payload(decode=True).decode()
                return main_text
    except Exception as e:
        raise Exception(f"Error processing EML file from GCS: {e}")
